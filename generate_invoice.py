import os
from docx import Document
from fpdf import FPDF

def count_words_docx(file_path):
    doc = Document(file_path)
    all_text = []

    all_text.extend([para.text.strip() for para in doc.paragraphs if para.text.strip()])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    all_text.append(cell_text)

    full_text = " ".join(all_text)
    words = full_text.split()
    return len(words)

def generate_invoice_data(folder_path, rate_per_word):
    total_words = 0
    total_cost = 0.0
    data = []

    docx_files = sorted([f for f in os.listdir(folder_path) if f.endswith(".docx")])
    for idx, filename in enumerate(docx_files, start=1):
        path = os.path.join(folder_path, filename)
        words = count_words_docx(path)
        price = round(words * rate_per_word, 1)
        data.append((idx, filename.replace('.docx', ''), words, rate_per_word, price))
        total_words += words
        total_cost += price

    return data, total_words, round(total_cost, 0)

def save_invoice_pdf(data, total_words, total_price, filename):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "", 10)

    col_widths = [10, 85, 20, 20, 25]
    headers = ["Sr.", "Discription", "Qty", "Rate", "Price"]

    def draw_row(values, bold=False):
        if bold:
            pdf.set_font("Arial", "B", 10)
        else:
            pdf.set_font("Arial", "", 10)

        for i, val in enumerate(values):
            align = "R" if i >= 2 else "L"
            pdf.cell(col_widths[i], 8, str(val), border=1, align=align)
        pdf.ln()

    # Header
    pdf.ln(5)
    draw_row(headers, bold=True)

    # Data rows
    for row in data:
        draw_row(row)

    # Total Rs
    pdf.set_font("Arial", "B", 10)
    pdf.cell(sum(col_widths[:4]), 8, "Total(Rs)", border=1, align="R")
    pdf.cell(col_widths[4], 8, str(total_price), border=1, align="R")
    pdf.ln()

    # Total Words
    pdf.cell(sum(col_widths[:4]), 8, "Total Words", border=1, align="R")
    pdf.cell(col_widths[4], 8, str(total_words), border=1, align="R")

    pdf.output(filename)
    print(f"\n✅ PDF invoice saved as: {filename}")

def main():
    folder = os.path.join(os.getcwd(), "articles")
    if not os.path.exists(folder):
        print("❌ Folder 'articles' not found.")
        return

    try:
        rate = float(input("Enter rate per word (e.g., 1.4): ").strip())
    except:
        print("❌ Invalid rate.")
        return

    output_name = input("Enter name for the PDF file (without extension): ").strip()
    if not output_name:
        output_name = "invoice"
    if not output_name.lower().endswith(".pdf"):
        output_name += ".pdf"

    data, total_words, total_price = generate_invoice_data(folder, rate)
    if not data:
        print("❌ No .docx files found.")
        return

    save_invoice_pdf(data, total_words, total_price, output_name)

if __name__ == "__main__":
    main()
